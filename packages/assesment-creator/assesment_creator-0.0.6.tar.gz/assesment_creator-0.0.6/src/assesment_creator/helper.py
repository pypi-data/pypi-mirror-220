from ensure import ensure_annotations
import docx
import os


@ensure_annotations
def output_file(file_name: str) -> str:
    output_file = ""
    splitted_file_name = file_name.split(".")
    print(splitted_file_name)

    if "." in file_name and splitted_file_name[-1] != "docx":
        output_file = splitted_file_name[0] + ".docx"

    elif "." not in file_name:
        output_file = file_name + ".docx"
    else:
        output_file = file_name
    return output_file


@ensure_annotations
def create_docx(qno_list: list, output_file: str = "assessment.docx") -> str:
    doc = docx.Document()
    # Add a Title to the document
    doc.add_heading("Assesment", 0)
    for i, qno in enumerate(qno_list):
        if len(qno[0]) != 0:
            doc.add_paragraph(str(qno[0]))
        else:
            doc.add_paragraph(str(qno[1]))
    os.makedirs("data", exist_ok=True)
    doc.save(os.path.join(os.path.curdir, "data", output_file))
    return "file created successfully"
